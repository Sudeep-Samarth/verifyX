"""
Smoke test for brd_parser.py — covers heading detection, role labelling,
full record building, no-objectives-heading BRD, and the new txt/docx paths.
Run: python backend/ingestion/test_brd_parser.py
"""
import sys
import os
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from brd_parser import (
    _assign_role,
    _is_heading,
    build_semantic_records,
    _blocks_from_txt,
    _blocks_from_docx,
    DEFAULT_ROLE,
)

PASS = "[PASS]"
FAIL = "[FAIL]"
failures = 0


def check(desc: str, got, expected):
    global failures
    if got == expected:
        print(f"{PASS}  {desc}")
    else:
        print(f"{FAIL}  {desc}")
        print(f"       Expected: {expected!r}")
        print(f"       Got:      {got!r}")
        failures += 1


# ------------------------------------------------------------------
# 1. Heading detection
# ------------------------------------------------------------------
check("Numbered heading detected",        _is_heading("2.1 Business Objectives"), True)
check("ALL-CAPS heading detected",        _is_heading("EXECUTIVE SUMMARY"),        True)
check("Long paragraph not a heading",     _is_heading("The system must support multi-user login and access control."), False)
check("Single word not a heading",        _is_heading("Overview"),                 False)

# ------------------------------------------------------------------
# 2. Role assignment from heading
# ------------------------------------------------------------------
check("Heading 'Business Objectives' -> Objective",
      _assign_role("2. Business Objectives", ""), "Objective")
check("Heading 'SCOPE OF WORK' -> Scope",
      _assign_role("SCOPE OF WORK", ""), "Scope")
check("Heading 'Risk Assessment' -> Risk",
      _assign_role("Risk Assessment", ""), "Risk")
check("Heading 'Timeline & Milestones' -> Timeline",
      _assign_role("Timeline & Milestones", ""), "Timeline")

# ------------------------------------------------------------------
# 3. Role assignment from body text (vague heading)
# ------------------------------------------------------------------
check("Body 'shall' -> Requirement",
      _assign_role("General Information", "The system shall support export to CSV and Excel."), "Requirement")
check("Body 'assumption' -> Assumption",
      _assign_role("Notes", "It is assumed that all users will have stable internet connectivity."), "Assumption")
check("Body 'stakeholder' -> Stakeholder",
      _assign_role("People", "The primary stakeholder is the Chief Operating Officer."), "Stakeholder")

# ------------------------------------------------------------------
# 4. Full record building (mock blocks)
# ------------------------------------------------------------------
mock_blocks = [
    {"page_number": 1, "text": "1. Executive Summary",   "is_heading": True},
    {"page_number": 1, "text": "This document provides an overview of the background and context of the project initiative.", "is_heading": False},
    {"page_number": 2, "text": "2. Business Objectives", "is_heading": True},
    {"page_number": 2, "text": "The primary goal is to reduce operational costs by thirty percent within twelve months.", "is_heading": False},
    {"page_number": 3, "text": "RISKS AND MITIGATION",   "is_heading": True},
    {"page_number": 3, "text": "There is a risk of scope creep due to unclear requirements from the client side.", "is_heading": False},
    {"page_number": 4, "text": "3. Scope of Work",       "is_heading": True},
    {"page_number": 4, "text": "The project covers integration with the existing ERP system and data migration.", "is_heading": False},
]
records  = build_semantic_records(mock_blocks)
role_map = {r["section_title"]: r["semantic_role"] for r in records}

check("Executive Summary section -> Background",   role_map.get("1. Executive Summary"),  "Background")
check("Business Objectives section -> Objective",  role_map.get("2. Business Objectives"), "Objective")
check("RISKS AND MITIGATION section -> Risk",      role_map.get("RISKS AND MITIGATION"),   "Risk")
check("Scope of Work section -> Scope",            role_map.get("3. Scope of Work"),        "Scope")

# ------------------------------------------------------------------
# 5. No-objectives-heading BRD still produces records
# ------------------------------------------------------------------
no_obj_blocks = [
    {"page_number": 1, "text": "PROJECT OVERVIEW",          "is_heading": True},
    {"page_number": 1, "text": "This initiative aims to digitise the procurement workflow for internal teams.", "is_heading": False},
    {"page_number": 2, "text": "FUNCTIONAL REQUIREMENTS",   "is_heading": True},
    {"page_number": 2, "text": "The system must support role-based access control with at least five different permission levels.", "is_heading": False},
]
no_obj_records = build_semantic_records(no_obj_blocks)
check("No-objectives BRD still produces records", len(no_obj_records) > 0, True)
check("Functional Requirements heading -> Requirement",
      next((r["semantic_role"] for r in no_obj_records if r["section_title"] == "FUNCTIONAL REQUIREMENTS"), None),
      "Requirement")

# ------------------------------------------------------------------
# 6. TXT extractor — write a temp file and verify blocks
# ------------------------------------------------------------------
TXT_CONTENT = """\
1. Executive Summary
This project aims to streamline the onboarding process for new clients joining the platform.

SCOPE OF WORK
The scope covers user registration, profile management, and document upload functionalities.

2.1 Business Objectives
The primary goal is to cut onboarding time from two weeks to three business days.
"""
with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as tmp:
    tmp.write(TXT_CONTENT)
    tmp_txt_path = tmp.name

try:
    txt_blocks = _blocks_from_txt(tmp_txt_path)
    heading_texts = [b["text"] for b in txt_blocks if b["is_heading"]]
    check("TXT extractor detects '1. Executive Summary' as heading",
          "1. Executive Summary" in heading_texts, True)
    check("TXT extractor detects 'SCOPE OF WORK' as heading",
          "SCOPE OF WORK" in heading_texts, True)
    check("TXT extractor detects '2.1 Business Objectives' as heading",
          "2.1 Business Objectives" in heading_texts, True)
    txt_records = build_semantic_records(txt_blocks)
    txt_role_map = {r["section_title"]: r["semantic_role"] for r in txt_records}
    check("TXT: Scope of Work role -> Scope",
          txt_role_map.get("SCOPE OF WORK"), "Scope")
    check("TXT: Business Objectives role -> Objective",
          txt_role_map.get("2.1 Business Objectives"), "Objective")
finally:
    os.unlink(tmp_txt_path)

# ------------------------------------------------------------------
# 7. DOCX extractor — write a temp .docx and verify blocks
# ------------------------------------------------------------------
try:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph("This is a background overview of the project context and its key drivers.")
    doc.add_heading("RISKS AND MITIGATION", level=2)
    doc.add_paragraph("There is a significant risk of budget overrun if scope is not tightly managed.")
    doc.add_heading("2.1 Functional Requirements", level=2)
    doc.add_paragraph("The system must support concurrent users with role-based permissions at all levels.")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_docx_path = tmp.name
    doc.save(tmp_docx_path)

    try:
        docx_blocks  = _blocks_from_docx(tmp_docx_path)
        docx_headings = [b["text"] for b in docx_blocks if b["is_heading"]]
        check("DOCX extractor detects 'Heading 1' paragraph as heading",
              "1. Project Overview" in docx_headings, True)
        check("DOCX extractor detects 'Heading 2' paragraph as heading",
              "RISKS AND MITIGATION" in docx_headings, True)
        docx_records  = build_semantic_records(docx_blocks)
        docx_role_map = {r["section_title"]: r["semantic_role"] for r in docx_records}
        check("DOCX: Risk section -> Risk",
              docx_role_map.get("RISKS AND MITIGATION"), "Risk")
        check("DOCX: Requirements section -> Requirement",
              docx_role_map.get("2.1 Functional Requirements"), "Requirement")
    finally:
        os.unlink(tmp_docx_path)

except ImportError:
    print("⚠️  python-docx not installed — skipping DOCX tests (run: pip install python-docx)")

# ------------------------------------------------------------------
print(f"\n{'All tests passed!' if failures == 0 else f'{failures} test(s) failed.'}")
sys.exit(failures)
